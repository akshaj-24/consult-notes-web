from __future__ import annotations

from io import BytesIO

from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Inches, Pt


def attach_docx_to_consult(consult_note):
    document = Document()
    style = consult_note.format_snapshot_json or {}
    section = document.sections[0]
    section.left_margin = Inches(float(style.get('left_margin', 1.0)))
    section.right_margin = Inches(float(style.get('right_margin', 1.0)))
    section.top_margin = Inches(float(style.get('top_margin', 1.0)))
    section.bottom_margin = Inches(float(style.get('bottom_margin', 1.0)))
    if style.get('header_text'):
        section.header.paragraphs[0].text = style['header_text']
    if style.get('footer_text'):
        section.footer.paragraphs[0].text = style['footer_text']

    heading_size = Pt(style.get('heading_font_size', 14))
    subheading_size = Pt(style.get('subheading_font_size', 12))
    body_size = Pt(style.get('body_font_size', 11))
    heading_font = style.get('heading_font_family', 'Arial')
    subheading_font = style.get('subheading_font_family', 'Arial')
    body_font = style.get('body_font_family', 'Arial')
    heading_bold = style.get('heading_bold', True)
    subheading_bold = style.get('subheading_bold', True)
    body_bold = style.get('body_bold', False)

    for index, raw_line in enumerate((consult_note.note_text or '').splitlines()):
        line = raw_line.strip()
        if not line:
            continue
        if index == 0:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            run.font.name = heading_font
            run.font.size = heading_size
            continue
        if line.startswith('-- '):
            paragraph = document.add_paragraph(style='List Bullet')
            run = paragraph.add_run(line[3:])
            run.font.name = body_font
            run.font.size = body_size
            run.bold = body_bold
            continue
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        if line.isupper() and len(line) <= 80:
            run.font.name = subheading_font
            run.font.size = subheading_size
            run.bold = subheading_bold
            if style.get('show_lines_between_sections'):
                paragraph.paragraph_format.space_before = Pt(8)
        else:
            run.font.name = body_font
            run.font.size = body_size
            run.bold = body_bold
        paragraph.paragraph_format.space_after = Pt(float(style.get('section_spacing', 1.0)) * 6)

    buffer = BytesIO()
    document.save(buffer)
    filename = f'{consult_note.patient.mrn}-{consult_note.pk}.docx'
    consult_note.docx_file.save(filename, ContentFile(buffer.getvalue()), save=False)
    consult_note.source_storage_path = consult_note.docx_file.name
    consult_note.save(update_fields=['docx_file', 'source_storage_path', 'updated_at'])
    return consult_note.docx_file
